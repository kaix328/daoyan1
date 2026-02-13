from io import BytesIO
from docx import Document
import pandas as pd

def convert_to_docx(script_text):
    doc = Document()
    # Clean markdown
    lines = script_text.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('|'):
            # Simple table handling (better use regex or pandas for complex tables, but this is basic)
            # For simplicity in this port, we just dump text if it's a table row
            # A full implementation would parse markdown tables.
            # Let's try to be a bit smarter: just add as paragraph for now to ensure robustness
            p = doc.add_paragraph(line)
            p.style.font.name = 'Courier New'
        else:
            doc.add_paragraph(line)
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def convert_to_xlsx(script_text):
    bio = BytesIO()
    # Create Excel writer
    try:
        writer = pd.ExcelWriter(bio, engine='openpyxl')
    except ImportError:
        # Fallback if openpyxl not installed, though unlikely if pandas is used
        print("[WARN] openpyxl not found, trying default")
        writer = pd.ExcelWriter(bio)

    lines = script_text.split('\n')
    tables = []
    current_table = []
    in_table = False
    
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('|'):
            in_table = True
            current_table.append(stripped)
        else:
            if in_table:
                tables.append(current_table)
                current_table = []
                in_table = False
    if in_table:
        tables.append(current_table)
        
    if not tables:
        # No tables found, just put text in A1
        df = pd.DataFrame({'Content': [script_text]})
        df.to_excel(writer, sheet_name='Script', index=False)
    else:
        for i, table_lines in enumerate(tables):
            try:
                data = []
                for row in table_lines:
                    # Check if it is a separator row (contains only dashes, pipes, spaces, colons)
                    if set(row.replace('|','').strip()) <= {'-', ':', ' '}:
                        continue
                        
                    # Split by pipe, assuming markdown table format
                    # Note: This is a simple split. Escaped pipes \| might break it, but sufficient for this context.
                    cells = [c.strip() for c in row.strip('|').split('|')]
                    data.append(cells)
                
                if data:
                    headers = data[0]
                    rows = data[1:]
                    
                    # Normalize row lengths
                    cleaned_rows = []
                    for r in rows:
                        if len(r) == len(headers):
                            cleaned_rows.append(r)
                        elif len(r) < len(headers):
                            cleaned_rows.append(r + [''] * (len(headers) - len(r)))
                        else:
                            cleaned_rows.append(r[:len(headers)])
                            
                    df = pd.DataFrame(cleaned_rows, columns=headers)
                    sheet_name = f'Table_{i+1}'
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
            except Exception as e:
                print(f"[ERROR] Error parsing table {i}: {e}")
                
    writer.close()
    bio.seek(0)
    return bio
