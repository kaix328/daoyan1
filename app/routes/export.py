"""
AI Director Assistant - 导出功能模块
"""

from flask import Blueprint, request, jsonify, current_app, send_file
from app.utils.helpers import Utils, ValidationError
from app.models.database import DatabaseManager
from datetime import datetime
import io
from docx import Document
import json

export_bp = Blueprint('export', __name__)

# 初始化数据库
db = DatabaseManager()

@export_bp.route('/pdf', methods=['POST'])
def export_pdf():
    """导出PDF文件"""
    try:
        data = request.get_json()
        if not data:
            return Utils.create_error_response("请求数据不能为空")
        
        # 验证必需字段
        if 'content' not in data:
            return Utils.create_error_response("缺少内容字段")
        
        content = data['content']
        title = data.get('title', 'AI Director Assistant 报告')
        
        # 创建PDF内容
        pdf_content = f"""
# {title}

生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

---

{content}

---

*本报告由 AI Director Assistant 自动生成
        """
        
        # 创建内存中的PDF文件
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # 构建PDF内容
        story = []
        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 12))
        story.append(Paragraph("---", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # 添加主要内容
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                story.append(Paragraph(line, styles['Normal']))
            else:
                story.append(Spacer(1, 6))
        
        story.append(Spacer(1, 12))
        story.append(Paragraph("---", styles['Normal']))
        story.append(Spacer(1, 12))
        story.append(Paragraph("*本报告由 AI Director Assistant 自动生成", styles['Italic']))
        
        doc.build(story)
        
        # 准备文件发送
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
            mimetype='application/pdf'
        )
        
    except ValidationError as e:
        return Utils.create_error_response(str(e))
    except Exception as e:
        Utils.log_error(e, "导出PDF")
        return Utils.create_error_response("PDF导出失败")

@export_bp.route('/word', methods=['POST'])
def export_word():
    """导出Word文档"""
    try:
        data = request.get_json()
        if not data:
            return Utils.create_error_response("请求数据不能为空")
        
        # 验证必需字段
        if 'content' not in data:
            return Utils.create_error_response("缺少内容字段")
        
        content = data['content']
        title = data.get('title', 'AI Director Assistant 报告')
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading(title, 0)
        
        # 添加生成时间
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 添加分隔线
        doc.add_paragraph("-" * 50)
        
        # 添加主要内容
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                if line.startswith('#'):
                    # 标题
                    level = min(line.count('#'), 6)
                    heading_text = line.replace('#', '').strip()
                    doc.add_heading(heading_text, level)
                elif line.startswith('*') or line.startswith('-'):
                    # 列表项
                    doc.add_paragraph(line.strip(), style='List Bullet')
                elif line.startswith('1.') or line.startswith('2.'):
                    # 编号列表
                    doc.add_paragraph(line.strip(), style='List Number')
                else:
                    # 普通段落
                    doc.add_paragraph(line.strip())
        
        # 添加页脚
        doc.add_paragraph("-" * 50)
        doc.add_paragraph("本报告由 AI Director Assistant 自动生成")
        
        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except ValidationError as e:
        return Utils.create_error_response(str(e))
    except Exception as e:
        Utils.log_error(e, "导出Word")
        return Utils.create_error_response("Word导出失败")

@export_bp.route('/json', methods=['POST'])
def export_json():
    """导出JSON文件"""
    try:
        data = request.get_json()
        if not data:
            return Utils.create_error_response("请求数据不能为空")
        
        # 验证必需字段
        if 'content' not in data:
            return Utils.create_error_response("缺少内容字段")
        
        content = data['content']
        title = data.get('title', 'AI Director Assistant 报告')
        
        # 构建JSON数据
        export_data = {
            'title': title,
            'generated_at': datetime.now().isoformat(),
            'content': content,
            'metadata': data.get('metadata', {})
        }
        
        # 创建JSON文件
        json_str = json.dumps(export_data, ensure_ascii=False, indent=2)
        buffer = io.BytesIO()
        buffer.write(json_str.encode('utf-8'))
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mimetype='application/json'
        )
        
    except ValidationError as e:
        return Utils.create_error_response(str(e))
    except Exception as e:
        Utils.log_error(e, "导出JSON")
        return Utils.create_error_response("JSON导出失败")

@export_bp.route('/script/<int:script_id>', methods=['GET'])
def export_script(script_id):
    """导出指定脚本"""
    try:
        # 获取脚本数据
        scripts = db.get_all_scripts(100)  # 获取所有脚本
        script = None
        
        for s in scripts:
            if s['id'] == script_id:
                script = s
                break
        
        if not script:
            return Utils.create_error_response("脚本不存在")
        
        # 构建导出格式
        format_type = request.args.get('format', 'json').lower()
        
        if format_type == 'json':
            buffer = io.BytesIO()
            json_data = json.dumps(script, ensure_ascii=False, indent=2)
            buffer.write(json_data.encode('utf-8'))
            buffer.seek(0)
            
            return send_file(
                buffer,
                as_attachment=True,
                download_name=f"script_{script_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mimetype='application/json'
            )
        
        elif format_type == 'txt':
            content = f"""
脚本ID: {script['id']}
主题: {script['theme']}
类型: {script['script_type']}
平台: {script['platform']}
创建时间: {script['created_at']}
收藏: {'是' if script['is_favorite'] else '否'}

内容:
{script['content']}

元数据:
{json.dumps(script.get('metadata', {}), ensure_ascii=False, indent=2)}
"""
            
            buffer = io.BytesIO()
            buffer.write(content.encode('utf-8'))
            buffer.seek(0)
            
            return send_file(
                buffer,
                as_attachment=True,
                download_name=f"script_{script_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mimetype='text/plain'
            )
        
        else:
            return Utils.create_error_response("不支持的导出格式")
        
    except ValidationError as e:
        return Utils.create_error_response(str(e))
    except Exception as e:
        Utils.log_error(e, "导出脚本")
        return Utils.create_error_response("导出失败")